"""Tests du renderer .docx ATS (M5/M6/M7) — oracle = le gabarit de référence.

Le gabarit n'est pas décrit de mémoire : il est LU (`cv/template/Template_Resume.docx`,
61 paragraphes, 0 table, style `Normal` seul, titres en petites capitales grasses,
entrées alignées par TABULATIONS) et sert d'oracle structurel à la sortie produite.

Contrainte héritée (ADR-002) : `test_phone_never_projected` interdit au téléphone de
traverser `build_structured_cv` (le dépôt est PUBLIC). L'overlay privé s'injecte donc
au seul point de rendu .docx — c'est ce que verrouillent les tests M7 ci-dessous.
"""
from __future__ import annotations

import collections
import copy
import functools
import io
import json
import pathlib
import sys
import types

import pytest
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Twips

import cv_docx
import cv_select

_REPO = pathlib.Path(__file__).resolve().parents[2]
_REFERENCE = _REPO / "cv" / "template" / "Template_Resume.docx"
_REAL_PROFILE = _REPO / "profile.json"
_PREFAB_FR = _REPO / "cv" / "prefab" / "full_fr.pdf"


# ── Fixtures ──────────────────────────────────────────────────────────────────

def _profile() -> dict:
    """Profil de test au SCHÉMA RÉEL de profile.json (location/division sur les
    expériences, projects typés, context = clé étrangère vers education/experiences)."""
    return {
        "$updated": "2026-07-28",
        "identity": {
            "first_name": "Robin", "last_name": "Denis", "email": "r@x.io",
            "phone": "+33 6 00 00 00 00",  # présent dans la source — jamais projeté
            "location": {"city": "Paris", "country": "France"},
            "links": {"linkedin": "https://www.linkedin.com/in/robin-x/",
                      "github": "https://github.com/RoJLD"},
            "tagline": {"fr": "Ingénierie Financière", "en": "Financial Engineering"},
        },
        "experiences": [
            {"id": "alten_2026", "company": "ALTEN",
             "division": {"fr": "Lab R&D", "en": "Lab R&D"},
             "location": "Sèvres, France",
             "title": {"fr": "Quant DeFi", "en": "DeFi Quant"},
             "start": "2026-02", "end": "2026-08", "current": True,
             "domains": ["quant"], "relevance": {"general": 0.9},
             "bullets": {"fr": ["Pilotage CryptoExploration"], "en": ["Led CryptoExploration"]}},
        ],
        "education": [
            {"id": "ece", "school": "ECE Paris",
             "title": {"fr": "ECE Paris, Cycle Ingénieur", "en": "ECE Paris, Engineering"},
             "org": {"fr": "Majeure IF&Q", "en": "Major FQE"}, "period": "2021 – 2026",
             "degree": "Diplôme d'Ingénieur",
             "courses_label": {"fr": "Cours clés", "en": "Key courses"},
             "courses": [{"fr": "Calcul stochastique", "en": "Stochastic calculus"}],
             "capstone": {"label": {"fr": "PFE (avec EY)", "en": "Capstone (EY)"},
                          "summary": {"fr": "Couverture dynamique", "en": "Dynamic hedging"}}},
        ],
        "projects": [
            {"id": "pfe-hedging", "type": "academic", "context": "ece", "date": "2025 — 2026",
             "name": {"fr": "Couverture dynamique (PFE × EY)", "en": "Dynamic hedging (Capstone × EY)"},
             "summary": {"fr": "Modélisation de stratégies de couverture.",
                         "en": "Modeling of hedging strategies."},
             "impact": {"fr": "Chef de projet", "en": "Project lead"},
             "domains": ["quant", "risk"], "stack": ["Python"]},
            {"id": "sudoku-cnn", "type": "academic", "context": "ece", "date": "ECE Paris",
             "name": {"fr": "Sudoku par CNN", "en": "Sudoku solving with CNN"},
             "summary": {"fr": "Résolution de grilles par réseau convolutif.",
                         "en": "Grid solving with a convolutional network."},
             "domains": ["ai"]},
            {"id": "elysium", "type": "personal", "context": "personal", "date": "2024 — présent",
             "name": {"fr": "ELYSIUM", "en": "ELYSIUM"},
             "summary": {"fr": "Plateforme cognitive.", "en": "Cognitive platform."},
             "domains": ["ai", "infra"]},
            {"id": "tms-bouygues", "type": "professional", "context": "bouygues_2025",
             "date": "2025", "name": {"fr": "TMS", "en": "TMS"},
             "summary": {"fr": "Treasury Management System.", "en": "Treasury Management System."},
             "domains": ["finance"]},
        ],
        "skills": {"programming": [{"name": "Python"}, {"name": "SQL"}]},
        "skills_labels": {"programming": {"fr": "Langages", "en": "Programming"}},
        "languages": [{"name": {"fr": "Anglais", "en": "English"},
                       "level": {"fr": "Courant (C1)", "en": "Fluent (C1)"}}],
        "certifications": ["TOEIC (925/990)"],
        "interests": [{"fr": "Échecs", "en": "Chess"}],
    }


_ATS_CFG = {"relevance_key": "general", "min_relevance": 0.0, "ats_extras": True}


def _scv(lang: str = "fr", cfg: dict | None = None, prof: dict | None = None) -> dict:
    prof = _profile() if prof is None else prof
    cfg = _ATS_CFG if cfg is None else cfg
    exps = cv_select.select_experiences(prof, cfg)
    return cv_select.build_structured_cv(prof, exps, lang, cfg)


def _text(doc) -> str:
    # délègue au helper exporté : sinon le module en expose un que rien n'exerce
    return cv_docx.document_text(doc)


def _prefab_cfgs() -> list[dict]:
    return json.loads((pathlib.Path(__file__).resolve().parent / "cv_profiles.json")
                      .read_text(encoding="utf-8"))["profiles"]


# ══ ORACLES INDÉPENDANTS ══════════════════════════════════════════════════════
# Une assertion qui compare la sortie à une constante de `cv_docx` ne teste que la
# plomberie : elle ne peut PAS voir la constante dériver. Mesuré 2026-07-28 —
# `RIGHT_TAB_EMU = 6570980 → 6350000`, `SZ_SECTION = Pt(10) → Pt(30)` et
# `_BULLET_INDENT = Pt(10) → Pt(40)` changeaient toutes trois le document rendu
# depuis le VRAI profile.json sans qu'aucun test ne bronche.
#
# Trois sources d'attendu INDÉPENDANTES de `cv_docx` sont donc utilisées ci-dessous :
#   1. `cv/template/Template_Resume.docx` — le gabarit versionné, LU (`_gabarit()`) ;
#   2. `profile.json` + la projection `cv_select` — la donnée réelle (`_expected_outline`) ;
#   3. `cv/prefab/full_fr.pdf` — un artefact SERVI, produit par l'AUTRE renderer.
#
# La « géométrie » suggérée comme oracle (largeur de page moins marges) est RÉFUTÉE
# par la mesure : le gabarit est en US Letter (7772400×10058400 EMU) et non en A4
# (7560000×10692000), et largeur − marges = 6742430 EMU, soit 171450 EMU À DROITE du
# taquet réel (6570980). Un test bâti sur cette formule serait rouge à tort.

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _all_runs_are(p, **attrs) -> bool:
    return bool(p.runs) and all(
        all(getattr(r, k, None) is v or getattr(r.font, k, None) is v for k, v in attrs.items())
        for r in p.runs)


def _numbering_hanging_twips(ref) -> int:
    """Retrait suspendu (twips) des listes à puces du gabarit, LU dans numbering.xml.

    Les puces du gabarit sont des listes Word (`numPr`) : leur retrait ne vit pas sur
    le paragraphe (`paragraph_format.first_line_indent` y vaut None, mesuré) mais dans
    la définition de numérotation. On la résout numId → abstractNum → niveau.
    """
    part = ref.part.numbering_part.element
    num2abs = {n.get(_W + "numId"): n.find(_W + "abstractNumId").get(_W + "val")
               for n in part.findall(_W + "num")}
    abs_lvls = {a.get(_W + "abstractNumId"): a.findall(_W + "lvl")
                for a in part.findall(_W + "abstractNum")}
    hangings = []
    for p in ref.paragraphs:
        numPr = p._p.find(_W + "pPr/" + _W + "numPr")
        if numPr is None or not p.text.strip():
            continue
        num_id = numPr.find(_W + "numId").get(_W + "val")
        ilvl = numPr.find(_W + "ilvl")
        ilvl = int(ilvl.get(_W + "val")) if ilvl is not None else 0
        lvls = abs_lvls.get(num2abs.get(num_id), [])
        if ilvl >= len(lvls):
            continue
        ind = lvls[ilvl].find(_W + "pPr/" + _W + "ind")
        if ind is not None and ind.get(_W + "hanging"):
            hangings.append(int(ind.get(_W + "hanging")))
    assert hangings, "aucune puce numérotée trouvée dans le gabarit de référence"
    return min(hangings)


@functools.lru_cache(maxsize=1)
def _gabarit() -> dict:
    """Toutes les mesures du gabarit de référence, LUES dans le fichier.

    Aucune valeur ne vient de `cv_docx` : c'est le `.docx` versionné qui répond.
    """
    ref = cv_docx.Document(str(_REFERENCE))
    sec, normal = ref.sections[0], ref.styles["Normal"]
    sizes = collections.Counter(r.font.size for p in ref.paragraphs for r in p.runs
                                if r.font.size is not None)
    # Titre de section du gabarit : paragraphe non vide, sans tabulation, dont TOUS
    # les runs sont gras + petites capitales (le critère « runs[0] » attraperait
    # aussi le §13, dont le premier run est une tabulation grasse — mesuré).
    titles = [p for p in ref.paragraphs
              if p.text.strip() and "\t" not in p.text
              and _all_runs_are(p, bold=True, small_caps=True)]
    title_sizes = {p.runs[0].font.size for p in titles}
    assert len(title_sizes) == 1, f"titres du gabarit à plusieurs corps : {title_sizes}"
    # Ligne d'entrée du gabarit = celle qui porte sa colonne DROITE (« City, Country »).
    entry_tabs = {(t.position, t.alignment)
                  for p in ref.paragraphs if "City, Country" in p.text
                  for t in p.paragraph_format.tab_stops
                  if t.alignment == WD_TAB_ALIGNMENT.RIGHT}
    assert len(entry_tabs) == 1, f"taquet droit des lignes d'entrée ambigu : {entry_tabs}"
    (tab_pos, tab_align), = entry_tabs
    return {
        "page": (sec.page_width, sec.page_height),
        "margins": (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin),
        "font": normal.font.name,
        "space_after": normal.paragraph_format.space_after,
        "line_spacing": normal.paragraph_format.line_spacing,
        "sz_body": sizes.most_common(1)[0][0],
        "sz_section": title_sizes.pop(),
        "sz_name": max(sizes),
        "tab_position": tab_pos,
        "tab_alignment": tab_align,
        "section_titles": {p.text.strip() for p in titles},
        "blank_paragraphs": sum(1 for p in ref.paragraphs if not p.text.strip()),
        "list_hanging": Twips(_numbering_hanging_twips(ref)),
    }


def _tabbed(left: str, right: str) -> str:
    return f"{left}\t{right}" if right else left


def _expected_outline(scv: dict, private: dict | None = None) -> list[str]:
    """MODÈLE-MONDE indépendant : les lignes que le `.docx` ATS doit porter.

    Calculé depuis la PROJECTION (`cv_select`) et l'overlay — c'est-à-dire depuis la
    DONNÉE, jamais depuis `cv_docx`. Le contrat du renderer est « toute valeur portée
    par la projection arrive dans le document, sur sa propre ligne, dans l'ordre de la
    projection » : ce modèle l'exprime, et l'égalité avec `[p.text for p in
    doc.paragraphs]` le vérifie dans les DEUX sens (rien ne manque, rien ne s'ajoute).

    Mesuré : sans lui, faire disparaître compétences + certifications + langues
    (62 → 55 paragraphes), le diplôme, les cours, le PFE, le sous-titre du header ou
    l'`impact` d'un projet laissait la suite VERTE.
    """
    lang = scv.get("lang") or "fr"
    lab = cv_docx._LABELS[lang]          # jeu de libellés gardé par ailleurs contre le gabarit
    priv = private or {}
    avail = priv.get("availability")
    avail = avail.get(lang) if isinstance(avail, dict) else avail
    idy = scv.get("identity") or {}

    lines = [idy.get("name") or ""]
    if idy.get("title"):
        lines.append(idy["title"])
    if avail:
        lines.append(str(avail))
    contact = " • ".join(x for x in (idy.get("location"), idy.get("email"),
                                     priv.get("phone"), idy.get("linkedin"),
                                     idy.get("github")) if x)
    if contact:
        lines.append(contact)

    exps = [s for s in scv.get("sections") or [] if (s or {}).get("kind") == "experience"]
    if exps:
        lines.append(lab["experience"])
        for s in exps:
            head = " — ".join(x for x in (s.get("company"), s.get("division")) if x)
            lines.append(_tabbed(head, s.get("location") or ""))
            lines.append(_tabbed(s.get("title") or "", s.get("dates") or ""))
            lines += ["• " + str(b) for b in s.get("bullets") or []]

    projects = scv.get("projects") or []
    if projects:
        lines.append(lab["projects"])
        for pr in projects:
            date, org = str(pr.get("date") or ""), str(pr.get("org") or "")
            lines.append(_tabbed(str(pr.get("name") or ""), date))
            if org and org.strip().casefold() != date.strip().casefold():
                lines.append(org)
            if pr.get("bullets"):
                lines += ["• " + str(b) for b in pr["bullets"]]
            elif pr.get("summary"):
                lines.append(str(pr["summary"]))
            if pr.get("impact"):
                lines.append(str(pr["impact"]))

    education = scv.get("education") or []
    if education:
        lines.append(lab["education"])
        for e in education:
            school, title = str(e.get("school") or ""), str(e.get("title") or "")
            lines.append(_tabbed(school, str(e.get("period") or "")))
            if title and school and title.startswith(school):
                title = ""      # « ECE Paris, Cycle Ingénieur » recouvre « ECE Paris »
            sub = " — ".join(x for x in (title, e.get("org")) if x)
            if sub:
                lines.append(sub)
            if e.get("degree"):
                lines.append(str(e["degree"]))
            if e.get("courses"):
                lines.append(f"{e.get('courses_label') or ''} : "
                             f"{' · '.join(str(c) for c in e['courses'])}".lstrip(" :"))
            cap = e.get("capstone")
            if isinstance(cap, dict):
                cap_txt = " — ".join(x for x in (cap.get("label"), cap.get("summary")) if x)
                if cap_txt:
                    lines.append(cap_txt)

    groups = scv.get("skills_groups") or []
    certs = scv.get("certifications") or []
    languages = scv.get("languages") or []
    if groups or certs or languages:
        lines.append(lab["skills"])
        for g in groups:
            if g.get("items"):
                lines.append(f"• {g.get('label') or ''} : "
                             f"{' · '.join(str(i) for i in g['items'])}")
        if certs:
            lines.append(f"• {lab['certifications']} : "
                         f"{' · '.join(str(c) for c in certs)}")
        lg_items = [" — ".join(x for x in (lg.get("name"), lg.get("level")) if x)
                    for lg in languages]
        lg_items = [x for x in lg_items if x]
        if lg_items:
            lines.append(f"• {lab['languages']} : {' · '.join(lg_items)}")

    interests = scv.get("interests") or []
    if interests:
        lines.append(lab["interests"])
        lines.append("• " + " · ".join(str(i) for i in interests))
    return lines


# ══ M6 — projection des projets ═══════════════════════════════════════════════

def test_select_projects_defaults_to_academic_only():
    got = [p["id"] for p in cv_select.select_projects(_profile(), {})]
    assert got == ["pfe-hedging", "sudoku-cnn"]  # ni personal ni professional


def test_select_projects_types_are_configurable():
    cfg = {"projects_types": ["personal", "professional"]}
    got = {p["id"] for p in cv_select.select_projects(_profile(), cfg)}
    assert got == {"elysium", "tms-bouygues"}


def test_select_projects_filters_by_domains():
    cfg = {"projects_types": ["academic", "personal"], "projects_domains_in": ["quant"]}
    got = [p["id"] for p in cv_select.select_projects(_profile(), cfg)]
    assert got == ["pfe-hedging"]  # sudoku (ai) et elysium (ai/infra) exclus


def test_select_projects_sorted_by_year_desc_then_id():
    """`date` est une chaîne libre ('2025 — 2026', 'ECE Paris') : le tri s'appuie
    sur la dernière année à 4 chiffres, 0 si aucune → déterministe et testable."""
    cfg = {"projects_types": ["academic", "personal", "professional"]}
    got = [p["id"] for p in cv_select.select_projects(_profile(), cfg)]
    #        2026            2025              2024        aucune année → 0, en dernier
    assert got == ["pfe-hedging", "tms-bouygues", "elysium", "sudoku-cnn"]
    assert cv_select._project_year("2025 — 2026") == 2026
    assert cv_select._project_year("2024 — présent") == 2024
    assert cv_select._project_year("ECE Paris") == 0


def test_select_projects_max_cap():
    cfg = {"projects_types": ["academic"], "max_projects": 1}
    assert len(cv_select.select_projects(_profile(), cfg)) == 1


def test_project_org_resolved_from_context_foreign_key():
    """`context` est une clé étrangère vers education[].id / experiences[].id —
    l'organisation affichée est LUE, jamais inventée."""
    prof = _profile()
    entries = cv_select.project_entries(
        prof, cv_select.select_projects(prof, {"projects_types": ["academic", "personal", "professional"]}), "fr")
    by_id = {e["id"]: e for e in entries}
    assert by_id["pfe-hedging"]["org"] == "ECE Paris"        # education.id == 'ece'
    assert by_id["elysium"]["org"] == ""                     # 'personal' n'est pas une clé
    assert by_id["tms-bouygues"]["org"] == ""                # 'bouygues_2025' absent du fixture
    # …et la résolution passe bien AUSSI par les expériences quand la clé existe :
    prof["experiences"].append({"id": "bouygues_2025", "company": "Bouygues Telecom",
                                "start": "2025-02", "relevance": {}, "domains": [],
                                "bullets": {"fr": [], "en": []}})
    again = cv_select.project_entries(
        prof, cv_select.select_projects(prof, {"projects_types": ["professional"]}), "fr")
    assert again[0]["org"] == "Bouygues Telecom"


def test_project_entries_resolve_bilingual_fields():
    prof = _profile()
    en = cv_select.project_entries(prof, cv_select.select_projects(prof, {}), "en")
    assert en[0]["name"] == "Dynamic hedging (Capstone × EY)"
    assert en[0]["summary"] == "Modeling of hedging strategies."
    # jamais de repr de dict dans un champ texte
    for e in en:
        for k in ("name", "date", "org", "summary", "impact"):
            assert "{" not in e[k], (e["id"], k, e[k])


def test_project_entries_flag_what_remains_to_be_written():
    """Aucun projet n'a de `bullets` : la projection doit le SIGNALER, pas fabriquer
    du contenu éditorial à la place de Robin."""
    prof = _profile()
    entries = cv_select.project_entries(prof, cv_select.select_projects(prof, {}), "fr")
    assert all(e["bullets"] == [] for e in entries)
    assert all(e["needs_bullets"] for e in entries)
    assert cv_select.projects_needing_bullets(entries) == ["pfe-hedging", "sudoku-cnn"]


def test_project_entries_use_bullets_when_they_exist():
    prof = _profile()
    prof["projects"][0]["bullets"] = {"fr": ["Modélisé le portefeuille"], "en": ["Modeled the book"]}
    entries = cv_select.project_entries(prof, cv_select.select_projects(prof, {}), "fr")
    assert entries[0]["bullets"] == ["Modélisé le portefeuille"]
    assert entries[0]["needs_bullets"] is False


# ══ Verrou de parité : la projection par défaut est INCHANGÉE ═════════════════

_HISTORICAL_SCV_KEYS = {
    "lang", "identity", "sections", "skills_groups", "skills_top",
    "education", "languages", "certifications", "interests", "footer",
}
_HISTORICAL_SECTION_KEYS = {"kind", "company", "title", "dates", "bullets"}


def test_default_structured_cv_shape_is_byte_for_byte_the_historical_one():
    """Le miroir JS (`assets/js/cv-select.js`) ne connaît PAS les champs ATS. Sans
    le drapeau `ats_extras`, la projection doit rester STRICTEMENT identique, sinon
    les 12 cas `scv:` du harnais de parité rougissent.

    Ce test est le garde-fou exprimé en Python : le harnais node ne tourne pas en CI
    (aucune CI sur ce dépôt — tâche J3), celui-ci tourne avec `pytest tools/`.
    """
    scv = _scv(cfg={"relevance_key": "general", "min_relevance": 0.0})
    assert set(scv) == _HISTORICAL_SCV_KEYS
    assert set(scv["sections"][0]) == _HISTORICAL_SECTION_KEYS


def test_prefab_cfgs_do_not_request_ats_extras():
    """`cv_profiles.json` alimente le harnais de parité (gen_parity_cases.py). Y poser
    `ats_extras` casserait la parité tant que `assets/js/cv-select.js` ne miroite pas
    les champs. Verrou explicite plutôt que découverte par un FAIL node."""
    cfgs = _prefab_cfgs()
    assert all(not c.get("ats_extras") for c in cfgs)
    assert all("projects_types" not in c for c in cfgs)


def _fake_chromium(monkeypatch) -> list[str]:
    """Chromium factice : capture le HTML injecté dans chaque PDF, sans navigateur.

    Le HTML passé à `page.set_content(...)` juste avant `page.pdf(...)` EST le PDF —
    c'est le seul endroit où « le gabarit est appliqué au PDF » veut dire quelque
    chose. On stube donc `playwright.sync_api`, PAS `build_cv_bank` : la ligne 46
    (`cv_render.render_html(scv)`) est réellement exécutée.
    """
    sink: list[str] = []

    class _Page:
        def set_content(self, html, **_kw):
            sink.append(html)

        def pdf(self, **kw):
            pathlib.Path(kw["path"]).write_bytes(b"%PDF-1.4 (factice)")

    class _Browser:
        def new_page(self):
            return _Page()

        def close(self):
            pass

    class _PW:
        chromium = types.SimpleNamespace(launch=lambda **_kw: _Browser())

        def __enter__(self):
            return self

        def __exit__(self, *_a):
            return False

    api = types.ModuleType("playwright.sync_api")
    api.sync_playwright = lambda: _PW()
    pkg = types.ModuleType("playwright")
    pkg.sync_api = api
    monkeypatch.setitem(sys.modules, "playwright", pkg)
    monkeypatch.setitem(sys.modules, "playwright.sync_api", api)
    return sink


def _prefab_html(cfgs: list[dict], out_dir: pathlib.Path, monkeypatch) -> list[str]:
    """HTML réellement injecté dans les PDF préfabriqués, cfg par cfg × langue."""
    import build_cv_bank
    sink = _fake_chromium(monkeypatch)
    monkeypatch.setattr(build_cv_bank, "_OUT", out_dir / "prefab")
    build_cv_bank.build(_profile(), cfgs)
    return sink


def test_a_template_declared_in_a_cfg_must_actually_change_the_pdf(tmp_path, monkeypatch):
    """J5 — anti-clé décorative, mesurée par le COMPORTEMENT et non par une signature.

    Version précédente : `not declared or accepts_template`, avec
    `accepts_template = len(signature(render_html).parameters) >= 2`. Mesuré : en
    simulant l'état POST-#18 (arité 2) AVEC un cfg portant `"template": "NIMPORTE_QUOI"`,
    elle passait — l'arité du renderer ne dit RIEN de ce que `build_cv_bank` lui
    transmet, encore moins de ce que le PDF contient. La formule devenait donc une
    TAUTOLOGIE au merge de #18 : elle ne pouvait plus jamais rougir, exactement
    quand elle devenait nécessaire.

    Ce qui compte : deux cfgs identiques au `template` près produisent-ils deux
    documents différents ? On l'exécute pour de vrai (vraie `build_cv_bank.build`,
    Chromium factice), puis on exige que « déclarer » et « être appliqué » aillent
    ENSEMBLE — la doctrine déjà écrite dans `$comment_template` de cv_profiles.json :
    ajouter la clé ici ET l'argument dans `build_cv_bank.py:46`, les deux ensemble.
    """
    base = {"id": "sonde", "label": {"fr": "Sonde", "en": "Probe"},
            "relevance_key": "general", "min_relevance": 0.0}
    a = _prefab_html([{**base, "template": "ats"}], tmp_path / "a", monkeypatch)
    b = _prefab_html([{**base, "template": "sobre"}], tmp_path / "b", monkeypatch)
    assert a and len(a) == len(b), "la sonde n'a produit aucun document"
    # …et ce HTML est bien celui d'un PDF ÉCRIT (pas d'un simple appel au renderer)
    assert sorted(p.name for p in (tmp_path / "a" / "prefab").glob("*.pdf")) == \
        ["sonde_en.pdf", "sonde_fr.pdf"]
    template_reaches_the_pdf = a != b

    declared = [c["id"] for c in _prefab_cfgs() if "template" in c]
    assert bool(declared) == template_reaches_the_pdf, (
        f"cfgs déclarant `template` : {declared or 'aucun'} ; changer le gabarit "
        f"change réellement le PDF : {template_reaches_the_pdf}. Les deux vont "
        f"ensemble — une clé que personne ne lit est décorative, une chaîne câblée "
        f"que personne ne déclare est morte.")


def test_ats_extras_adds_projects_and_entry_metadata():
    scv = _scv("fr")
    assert set(scv) == _HISTORICAL_SCV_KEYS | {"projects"}
    assert set(scv["sections"][0]) == _HISTORICAL_SECTION_KEYS | {"location", "division"}
    assert scv["sections"][0]["location"] == "Sèvres, France"
    assert scv["sections"][0]["division"] == "Lab R&D"
    assert [p["id"] for p in scv["projects"]] == ["pfe-hedging", "sudoku-cnn"]


def test_ats_extras_never_projects_the_phone():
    """M7 : même en mode ATS, le téléphone ne traverse JAMAIS la projection."""
    scv = _scv("fr")
    assert "phone" not in scv["identity"]
    assert "+33" not in json.dumps(scv, ensure_ascii=False)


# ══ M5 — le renderer .docx ════════════════════════════════════════════════════

def test_reference_gabarit_is_present_and_measurable():
    """Oracle lu, pas décrit. Si le gabarit disparaît, ce test tombe fort."""
    assert _REFERENCE.exists(), f"gabarit de référence absent : {_REFERENCE}"
    ref = cv_docx.Document(str(_REFERENCE))
    assert len(ref.paragraphs) == 61
    assert len(ref.tables) == 0
    assert {p.style.name for p in ref.paragraphs} == {"Normal"}


def test_output_shares_the_reference_structural_invariants():
    ref = cv_docx.Document(str(_REFERENCE))
    out = cv_docx.build_docx(_scv("fr"))
    for name, doc in (("référence", ref), ("sortie", out)):
        assert len(doc.tables) == 0, name
        assert {p.style.name for p in doc.paragraphs} == {"Normal"}, name
        titles = [p for p in doc.paragraphs
                  if p.runs and p.runs[0].font.small_caps and p.runs[0].bold]
        assert titles, f"aucun titre en petites capitales grasses : {name}"


@pytest.mark.parametrize("lang", ["fr", "en"])
def test_section_titles_themselves_carry_the_small_caps_of_the_gabarit(lang):
    """Propriété n°1 du gabarit ATS, verrouillée SUR LES TITRES et non « quelque part ».

    Formulation précédente : « il existe un paragraphe dont runs[0] est gras +
    petites capitales ». Mesuré : `_entry_line` en pose AUSSI (ligne d'entreprise),
    donc l'existence restait satisfaite ailleurs quand `_section_title` perdait ses
    deux attributs — muter `bold=True, small_caps=True` en `False, False` laissait
    la suite VERTE. Ici chaque titre est retrouvé par SON texte, puis ses attributs
    sont exigés un par un (une mutation partielle rougit aussi).
    """
    lab = cv_docx._LABELS[lang]
    doc = cv_docx.build_docx(_scv(lang))
    by_text = {p.text: p for p in doc.paragraphs}
    for key in ("experience", "projects", "education", "skills", "interests"):
        label = lab[key]
        p = by_text.get(label)
        assert p is not None, f"titre de section absent du document : {label!r}"
        assert len(p.runs) == 1, f"{label!r} : {len(p.runs)} runs, 1 attendu"
        run = p.runs[0]
        assert run.bold is True, f"{label!r} : titre de section non gras"
        assert run.font.small_caps is True, f"{label!r} : titre sans petites capitales"
        # Corps du titre : celui MESURÉ dans le gabarit, jamais `cv_docx.SZ_SECTION`
        # (mesuré : cette comparaison-là laissait passer Pt(10) → Pt(30)).
        assert run.font.size == _gabarit()["sz_section"], f"{label!r} : corps de titre"
    # …et la propriété n'est PAS satisfaite par n'importe quel paragraphe : la
    # prose et les puces restent en corps normal, sans petites capitales.
    prose = by_text["Modélisation de stratégies de couverture."
                    if lang == "fr" else "Modeling of hedging strategies."]
    assert prose.runs[0].font.small_caps is not True
    assert prose.runs[0].bold is not True


def test_page_setup_matches_the_shipped_pdfs():
    """Les 3 PDF réellement envoyés font 612x792 pt = US Letter, comme le gabarit.

    Les 4 marges sont comparées au gabarit LU (les marges haut/bas → 0 laissaient la
    suite verte : elles n'étaient pas gardées).
    """
    ref = _gabarit()
    doc = cv_docx.build_docx(_scv("fr"))
    s = doc.sections[0]
    assert (s.page_width, s.page_height) == (7772400, 10058400) == ref["page"]
    assert (s.left_margin, s.right_margin, s.top_margin, s.bottom_margin) == ref["margins"]
    assert doc.styles["Normal"].font.name == "Times New Roman" == ref["font"]


def test_the_gabarit_constants_are_the_measurements_of_the_reference_file():
    """Les constantes « MESURÉES » de `cv_docx` doivent l'être VRAIMENT.

    Chacune est confrontée à sa valeur LUE dans `cv/template/Template_Resume.docx`.
    Une constante qui dérive rougit ici, quel que soit le reste de la suite — c'est le
    contre-poison direct des trois assertions tautologiques mesurées le 2026-07-28.
    Attention au piège d'arrondi : python-docx stocke un taquet en TWIPS, donc muter
    `RIGHT_TAB_EMU` vers une valeur non multiple de 635 rougirait pour la mauvaise
    raison. Ici la comparaison porte sur la constante elle-même : pas d'aller-retour.
    """
    ref = _gabarit()
    assert cv_docx.RIGHT_TAB_EMU == ref["tab_position"]
    assert cv_docx.SZ_SECTION == ref["sz_section"]
    assert cv_docx.SZ_BODY == ref["sz_body"]
    assert cv_docx.SZ_NAME == ref["sz_name"]
    assert cv_docx.BASE_FONT == ref["font"]
    assert (cv_docx.PAGE_WIDTH_EMU, cv_docx.PAGE_HEIGHT_EMU) == ref["page"]
    assert (cv_docx.MARGIN_LEFT_EMU, cv_docx.MARGIN_RIGHT_EMU,
            cv_docx.MARGIN_TOP_EMU, cv_docx.MARGIN_BOTTOM_EMU) == ref["margins"]
    # …et la hiérarchie typographique qui n'est PAS dans le gabarit (le sous-titre
    # n'y existe pas) reste bornée par ce qui y est : corps < sous-titre < nom.
    assert ref["sz_body"] < cv_docx.SZ_TAGLINE < ref["sz_name"]


def test_english_section_labels_are_the_ones_written_in_the_reference_gabarit():
    """Les libellés EN ne sont pas une invention : ce sont les titres du gabarit.

    Cela ancre `_LABELS['en']` sur le fichier de référence — et donc le modèle-monde
    `_expected_outline`, qui s'en sert pour nommer les sections attendues.
    """
    en = cv_docx._LABELS["en"]
    assert _gabarit()["section_titles"] == {
        en[k] for k in ("experience", "projects", "education", "skills", "interests")}


def test_every_run_wears_the_font_of_the_gabarit():
    """La propriété n°1 du gabarit, gardée LÀ OÙ WORD LA LIT : sur chaque run.

    Mesuré : le style `Normal` était gardé, mais `_run` repose la police sur CHAQUE
    run et c'est le run qui l'emporte au rendu — `r.font.name = BASE_FONT →
    "Comic Sans MS"` sortait un CV intégralement en Comic Sans MS, suite VERTE.
    """
    ref_font = _gabarit()["font"]
    for lang in ("fr", "en"):
        doc = cv_docx.build_docx(_scv(lang), private={"phone": "+00 0",
                                                      "availability": {lang: "Dispo"}})
        runs = [r for p in doc.paragraphs for r in p.runs]
        assert runs, "document sans run"
        wrong = {r.font.name for r in runs} - {ref_font}
        assert not wrong, f"[{lang}] runs hors gabarit : {wrong}"
        assert doc.styles["Normal"].font.name == ref_font


def test_body_runs_and_default_style_carry_the_body_size_of_the_gabarit():
    """Corps 9 pt : sur les runs ET sur le style de repli.

    Mesuré : `SZ_BODY 9 → 14 pt` et `normal.font.size → Pt(20)` laissaient la suite
    verte. La taille du style compte même si tous les runs la reposent : c'est ce
    qu'héritera le prochain run écrit sans taille.
    """
    ref = _gabarit()
    doc = cv_docx.build_docx(_scv("fr"))
    assert doc.styles["Normal"].font.size == ref["sz_body"]
    body = [r for p in doc.paragraphs for r in p.runs if p.text.startswith("• ")]
    assert body, "aucune ligne de corps"
    assert {r.font.size for r in body} == {ref["sz_body"]}


def test_header_carries_the_type_scale_and_the_emphasis_of_the_gabarit():
    """Nom au plus grand corps du gabarit ET gras ; l'italique reste réservé.

    Mesuré : `SZ_NAME 20 → 9 pt`, le nom non gras, et la ligne de disponibilité non
    italique laissaient tous trois la suite verte. L'italique n'est pas décoratif :
    c'est ce qui distingue la phrase de disponibilité des trois autres lignes du
    header, toutes en romain.
    """
    ref = _gabarit()
    private = {"phone": "+00 0 00", "availability": {"fr": "Stage 6 mois"}}
    doc = cv_docx.build_docx(_scv("fr"), private=private)
    name, tagline, avail, contact = doc.paragraphs[:4]
    assert name.text == "Robin Denis"
    assert name.runs[0].font.size == ref["sz_name"], "le nom n'est pas au corps du gabarit"
    assert name.runs[0].bold is True, "le nom du header n'est pas gras"
    assert avail.text == "Stage 6 mois"
    assert avail.runs[0].italic is True, "la disponibilité doit être en italique"
    assert [p.runs[0].italic for p in (name, tagline, contact)] == [False, False, False]
    assert ref["sz_body"] < tagline.runs[0].font.size < ref["sz_name"], "échelle du header"


def test_paragraph_spacing_follows_the_normal_style_of_the_gabarit():
    """Interligne simple + aucun espace après : LU dans le style `Normal` du gabarit.

    Mesuré : `space_after Pt(0) → Pt(12)` et `line_spacing 1.0 → 2.0` étaient posés
    par `_para` sur chaque paragraphe et n'étaient gardés nulle part — un CV à
    interligne double déborde de la page sans qu'un test bronche.
    """
    ref = _gabarit()
    doc = cv_docx.build_docx(_scv("fr"))
    for p in doc.paragraphs:
        pf = p.paragraph_format
        assert pf.space_after == ref["space_after"], p.text
        assert pf.line_spacing == ref["line_spacing"], p.text


def test_section_titles_are_aired_by_space_before_because_the_output_has_no_blank_line():
    """Le gabarit aère ses sections par des paragraphes VIDES ; nous, non.

    Mesuré sur `Template_Resume.docx` : il compte des paragraphes vides (aération
    Word). Un document ATS ne doit pas en porter — une ligne vide est du bruit à
    l'extraction. L'aération doit donc venir d'un `space_before` non nul sur les
    titres : le supprimer (`space_before=Pt(6) → None`) collait les sections au bloc
    précédent, suite VERTE.
    """
    ref = _gabarit()
    assert ref["blank_paragraphs"] > 0, "le gabarit n'aère pas par lignes vides ?"
    scv = _scv("fr")
    doc = cv_docx.build_docx(scv)
    assert not [p for p in doc.paragraphs if not p.text.strip()], "ligne vide dans le CV"
    labels = set(cv_docx._LABELS["fr"].values())
    titles = [p for p in doc.paragraphs if p.text in labels]
    assert len(titles) >= 4, [p.text for p in titles]
    for p in titles:
        assert (p.paragraph_format.space_before or 0) > 0, f"titre collé : {p.text!r}"

    # Même raisonnement UN CRAN PLUS BAS : sans ligne vide, deux entrées consécutives
    # se touchent si leur première ligne n'est pas aérée (mesuré : retirer
    # `space_before=Pt(3)` des lignes d'entrée laissait la suite verte). Les têtes
    # d'entrée sont calculées depuis la DONNÉE, pas relevées dans le document.
    heads = {_tabbed(" — ".join(x for x in (s.get("company"), s.get("division")) if x),
                     s.get("location") or "") for s in scv["sections"]}
    heads |= {_tabbed(p.get("name") or "", p.get("date") or "") for p in scv["projects"]}
    heads |= {_tabbed(e.get("school") or "", e.get("period") or "")
              for e in scv["education"]}
    seen = [p for p in doc.paragraphs if p.text in heads]
    assert len(seen) == len(heads), ([p.text for p in seen], heads)
    for p in seen:
        assert (p.paragraph_format.space_before or 0) > 0, f"entrée collée : {p.text!r}"
    # …et l'aération est HIÉRARCHIQUE : une section respire plus qu'une entrée.
    assert min(p.paragraph_format.space_before for p in titles) >= \
        max(p.paragraph_format.space_before for p in seen)


def test_a_section_of_another_kind_is_not_printed_under_the_experience_heading():
    """`sections[].kind` n'est pas décoratif : c'est le filtre du bloc « Expériences ».

    Le schéma `sections` est partagé avec le miroir JS et porte un `kind` justement
    parce que d'autres natures peuvent y vivre. Mesuré : relâcher le filtre
    (`if s_` au lieu de `if s_.get("kind") == "experience"`) laissait la suite verte —
    la donnée d'aujourd'hui ne contient que des expériences, donc aucun test ne
    pouvait voir la différence. On fabrique donc la section qui manque.
    """
    scv = copy.deepcopy(_scv("fr"))
    scv["sections"].append({"kind": "volunteering", "company": "Croix-Rouge",
                            "division": "", "location": "Paris",
                            "title": "Bénévole", "dates": "2023",
                            "bullets": ["Maraudes hebdomadaires"]})
    texts = [p.text for p in cv_docx.build_docx(scv).paragraphs]
    intrus = [t for t in texts if "Croix-Rouge" in t or "Maraudes" in t]
    assert not intrus, f"section `volunteering` rendue sous « Expériences » : {intrus}"
    assert texts == _expected_outline(scv)


def test_entry_lines_are_aligned_by_tabulations_with_a_right_stop():
    """Position ET alignement du taquet, tous deux LUS dans le gabarit.

    Version précédente : `[t.position] == [cv_docx.RIGHT_TAB_EMU]` — la constante
    comparée à elle-même. Mesuré : `RIGHT_TAB_EMU → 6350000` (puis 5715000) laissait
    la suite VERTE alors que le taquet du document rendu bougeait de 2,4 mm ; et
    l'ALIGNEMENT n'était pas gardé du tout (`RIGHT → LEFT` : VERT), alors que
    « entrées alignées par tabulations avec un taquet DROIT » est LA mesure du gabarit.
    Ici l'attendu vient de `Template_Resume.docx` : les deux mutations rougissent.
    """
    ref = _gabarit()
    doc = cv_docx.build_docx(_scv("fr"))
    tabbed = [p for p in doc.paragraphs if "\t" in p.text]
    assert tabbed, "aucune ligne d'entrée tabulée"
    for p in tabbed:
        stops = [(t.position, t.alignment) for t in p.paragraph_format.tab_stops]
        assert stops == [(ref["tab_position"], ref["tab_alignment"])], p.text
    assert [p.text for p in tabbed][0] == "ALTEN — Lab R&D\tSèvres, France"


def test_bullets_are_literal_so_they_survive_text_extraction():
    """Les puces `numPr` du gabarit sont INVISIBLES à l'extraction de texte
    (mesuré : le §18 de la référence rend 'Executed', sans marqueur). Une puce
    littérale traverse tous les extracteurs — c'est le sens d'« ATS-safe »."""
    ref = cv_docx.Document(str(_REFERENCE))
    assert ref.paragraphs[18].text == "Executed"  # marqueur perdu côté référence
    doc = cv_docx.build_docx(_scv("fr"))
    assert any(p.text.startswith("• Pilotage CryptoExploration") for p in doc.paragraphs)


def test_sections_are_present_and_localized():
    fr = _text(cv_docx.build_docx(_scv("fr")))
    en = _text(cv_docx.build_docx(_scv("en")))
    assert "Expériences professionnelles" in fr and "Projets académiques" in fr
    assert "Professional Experience" in en and "Academic Projects" in en
    assert "Formation" in fr and "Education" in en


def test_projects_fall_back_to_prose_when_no_bullets_exist():
    """Le gabarit veut des puces menées par un verbe ; aucun projet n'en a. On rend
    la prose existante (`summary`) — jamais une puce fabriquée."""
    doc = cv_docx.build_docx(_scv("fr"))
    txt = _text(doc)
    assert "Modélisation de stratégies de couverture." in txt
    assert "Couverture dynamique (PFE × EY)" in txt


def test_project_bullets_are_rendered_when_the_data_carries_them():
    """Couverture du RENDU, pas seulement de la projection.

    Mesuré : le test M6 existant n'exerçait que `project_entries` ; muter
    `if pr.get("bullets"):` en `if False:` (cv_docx.py:340-342) laissait la suite
    verte, et `--cov` confirmait que les lignes 341-342 n'étaient JAMAIS exécutées.
    """
    prof = _profile()
    prof["projects"][0]["bullets"] = {
        "fr": ["Modélisé le portefeuille de couverture", "Backtesté la stratégie"],
        "en": ["Modeled the hedging book", "Backtested the strategy"]}
    doc = cv_docx.build_docx(_scv("fr", prof=prof))
    texts = [p.text for p in doc.paragraphs]
    assert "• Modélisé le portefeuille de couverture" in texts, texts
    assert "• Backtesté la stratégie" in texts, texts
    # une VRAIE ligne à puce (retrait négatif), pas une ligne de prose. L'attendu ne
    # peut PAS être `-cv_docx._BULLET_INDENT` : cette comparaison-là laissait passer
    # Pt(10) → Pt(40) (mesuré). Le gabarit fixe la borne — ses propres puces (listes
    # `numPr`) suspendent de 288 twips au niveau 0, lu dans numbering.xml.
    puce = doc.paragraphs[texts.index("• Backtesté la stratégie")]
    pf = puce.paragraph_format
    hanging = -pf.first_line_indent
    assert hanging > 0, "la puce ne suspend pas : le texte replié passerait sous le •"
    assert hanging <= _gabarit()["list_hanging"], (
        f"retrait suspendu {hanging} EMU > celui du gabarit "
        f"({_gabarit()['list_hanging']} EMU) : la puce décrocherait du corps du texte")
    assert pf.left_indent == hanging, "le • doit tomber sur la marge, pas au-delà"
    # …et le repli en prose ne s'ajoute PAS : les deux branches sont exclusives
    assert "Modélisation de stratégies de couverture." not in texts


def test_project_org_line_is_rendered_under_its_entry_line():
    """`org` vient de la clé étrangère `context` (découverte M6) — et doit ARRIVER
    dans le document. Mesuré : muter `if pr.get("org"):` (cv_docx.py:337) en
    `if False and pr.get("org"):` laissait la suite verte ; seule la projection
    était testée, jamais la LIGNE.
    """
    doc = cv_docx.build_docx(_scv("fr"))
    texts = [p.text for p in doc.paragraphs]
    entry = "Couverture dynamique (PFE × EY)\t2025 — 2026"
    assert entry in texts, texts
    i = texts.index(entry)
    assert texts[i + 1] == "ECE Paris", texts[i:i + 3]
    run = doc.paragraphs[i + 1].runs[0]
    assert run.italic is True and run.bold is False


@pytest.mark.parametrize("lang", ["fr", "en"])
def test_the_document_carries_exactly_the_lines_the_projection_provides(lang):
    """Le contrat du renderer, vérifié dans les DEUX sens contre le modèle-monde.

    Mesuré : faire disparaître en silence les compétences groupées + certifications +
    langues (62 → 55 paragraphes), le diplôme, les cours, le PFE, le sous-titre du
    header ou l'`impact` d'un projet laissait la suite VERTE — aucun test ne comptait
    ce que le document DEVAIT porter. L'attendu vient ici de la projection, pas du
    renderer : c'est la donnée qui dit combien de lignes il doit y avoir.
    """
    private = {"phone": "+00 0 00 00 00",
               "availability": {"fr": "Stage 6 mois", "en": "6-month internship"}}
    scv = _scv(lang)
    doc = cv_docx.build_docx(scv, private=private)
    assert [p.text for p in doc.paragraphs] == _expected_outline(scv, private)


def test_an_entry_and_its_sub_line_are_distinguished_the_way_the_gabarit_does_it():
    """Deux lignes par entrée, typographiquement DISTINCTES — comme dans le gabarit.

    Lu dans `Template_Resume.docx` : sur une ligne d'entrée le run gras est en petites
    capitales, sur la ligne juste dessous il ne l'est pas. Notre rendu reprend cette
    distinction et ajoute l'italique sur la sous-ligne (mesuré : la retirer laissait
    la suite verte, alors que le poste et les dates se confondaient alors avec
    l'entreprise).
    """
    ref = cv_docx.Document(str(_REFERENCE))
    idx = [i for i, p in enumerate(ref.paragraphs) if "City, Country" in p.text]
    assert idx, "le gabarit ne porte plus de ligne d'entrée"
    head_sc = {r.font.small_caps for i in idx for r in ref.paragraphs[i].runs
               if r.bold and r.text.strip()}
    sub_sc = {r.font.small_caps for i in idx if i + 1 < len(ref.paragraphs)
              for r in ref.paragraphs[i + 1].runs if r.bold and r.text.strip()}
    assert head_sc == {True} and True not in sub_sc, (head_sc, sub_sc)

    scv = _scv("fr")
    doc = cv_docx.build_docx(scv)
    texts = [p.text for p in doc.paragraphs]
    i = texts.index("ALTEN — Lab R&D\tSèvres, France")
    head, sub = doc.paragraphs[i].runs[0], doc.paragraphs[i + 1].runs[0]
    assert sub.text == scv["sections"][0]["title"], sub.text
    assert head.font.small_caps is True and sub.font.small_caps is False
    assert head.italic is False and sub.italic is True, "sous-ligne d'entrée non italique"


def test_the_contact_line_is_the_one_the_shipped_prefab_pdf_carries():
    """Oracle EXTERNE : `cv/prefab/full_fr.pdf`, servi par le bouton du site.

    Ce PDF est produit par l'AUTRE renderer (HTML → Chromium) depuis le MÊME
    profile.json. Sa ligne de contact est donc un attendu qu'aucune mutation de
    `cv_docx` ne peut suivre. Mesuré : retirer l'e-mail, ou linkedin+github, de la
    ligne de contact laissait la suite VERTE — or « Ville • email • téléphone •
    linkedin » est la composition des CV réellement envoyés.
    """
    pypdf = pytest.importorskip("pypdf", reason="lecture du PDF préfabriqué servi")
    if not _PREFAB_FR.exists():
        pytest.skip(f"préfabriqué absent : {_PREFAB_FR}")
    page = pypdf.PdfReader(str(_PREFAB_FR)).pages[0].extract_text()
    shipped = [ln.strip() for ln in page.splitlines() if "@" in ln]
    assert len(shipped) == 1, f"ligne de contact du PDF servi introuvable : {shipped}"

    prof = json.loads(_REAL_PROFILE.read_text(encoding="utf-8"))
    cfg = {"relevance_key": "general", "min_relevance": 0.0, "ats_extras": True}
    scv = cv_select.build_structured_cv(
        prof, cv_select.select_experiences(prof, cfg), "fr", cfg)
    texts = [p.text for p in cv_docx.build_docx(scv).paragraphs]
    assert shipped[0] in texts, (shipped[0], texts[:5])
    # …et le téléphone privé s'insère À SA PLACE, entre l'e-mail et linkedin (ADR-007 :
    # absent des artefacts publics, présent dans les documents adressés).
    sentinel = "SENTINELLE-TEL"
    with_phone = [p.text for p in cv_docx.build_docx(
        scv, private={"phone": sentinel}).paragraphs]
    line = next(t for t in with_phone if scv["identity"]["email"] in t)
    parts = line.split(" • ")
    assert parts.index(sentinel) == parts.index(scv["identity"]["email"]) + 1, parts
    assert line.replace(f" • {sentinel}", "") == shipped[0]


def _assert_no_right_column_echo(doc, where: str) -> None:
    """Aucune ligne ne reprend telle quelle la colonne DROITE de la ligne du dessus."""
    texts = [p.text for p in doc.paragraphs]
    for cur, nxt in zip(texts, texts[1:]):
        if "\t" not in cur:
            continue
        right = cur.split("\t", 1)[1].strip().casefold()
        assert not right or right != nxt.strip().casefold(), (
            f"[{where}] colonne droite répétée sur la ligne suivante : "
            f"{cur!r} puis {nxt!r}")


def test_the_org_line_never_echoes_the_date_column_above_it():
    """Défaut MESURÉ dans l'artefact livré : `project.date` vaut littéralement
    « ECE Paris » pour 8 des 10 projets académiques du vrai profile.json — soit
    exactement l'`org` résolu depuis `context`. Le rendu imprimait donc la même
    chaîne sur deux lignes consécutives du CV envoyé. Le fixture reproduit le
    quirk (`sudoku-cnn` : date == org == 'ECE Paris').
    """
    doc = cv_docx.build_docx(_scv("fr"))
    texts = [p.text for p in doc.paragraphs]
    i = texts.index("Sudoku par CNN\tECE Paris")
    assert texts[i + 1] != "ECE Paris", texts[i:i + 3]
    assert texts[i + 1].startswith("Résolution de grilles"), texts[i:i + 3]
    for lang in ("fr", "en"):
        _assert_no_right_column_echo(cv_docx.build_docx(_scv(lang)), f"fixture/{lang}")
    # l'org RESTE rendu quand il n'est pas un doublon (cf. pfe-hedging plus haut)
    assert "ECE Paris" in texts


def test_the_anti_echo_rule_suppresses_an_exact_repetition_and_nothing_more():
    """La PRÉCISION de la garde anti-doublon, dans ses deux sens.

    Mesuré : muter `!=` en `not in` laissait la suite VERTE — et faisait disparaître
    la ligne org dès que l'org est une SOUS-CHAÎNE de la date (`date='ECE Paris 2024'`,
    `org='ECE Paris'`). C'est le défaut SYMÉTRIQUE de celui que la garde répare : la
    sur-suppression silencieuse. Dans l'autre sens, la comparaison doit rester
    insensible à la casse : `date` est du texte libre saisi à la main.
    """
    prof = copy.deepcopy(_profile())
    prof["projects"][1]["date"] = "ECE Paris 2024"      # org 'ECE Paris' = préfixe strict
    texts = [p.text for p in cv_docx.build_docx(_scv("fr", prof=prof)).paragraphs]
    i = texts.index("Sudoku par CNN\tECE Paris 2024")
    assert texts[i + 1] == "ECE Paris", texts[i:i + 3]

    prof["projects"][1]["date"] = "ece paris"           # même chaîne, casse différente
    texts = [p.text for p in cv_docx.build_docx(_scv("fr", prof=prof)).paragraphs]
    i = texts.index("Sudoku par CNN\tece paris")
    assert texts[i + 1].startswith("Résolution de grilles"), texts[i:i + 3]


def test_document_text_keeps_one_line_per_paragraph():
    """`document_text` est ce que « voit un ATS » : une ligne par paragraphe.

    Mesuré : joindre par une espace au lieu de `\\n` laissait la suite verte — le CV
    devenait un pavé d'un seul tenant pour tout extracteur qui lit ligne à ligne.
    """
    doc = cv_docx.build_docx(_scv("fr"))
    txt = cv_docx.document_text(doc)
    assert txt.count("\n") == len(doc.paragraphs) - 1
    assert txt.splitlines() == [p.text for p in doc.paragraphs]


def test_renders_a_projection_built_without_ats_extras():
    """Le renderer doit accepter un `structured_cv` HISTORIQUE (sans `projects`,
    sans `location`/`division`) : c'est la forme que produit tout appelant existant,
    et c'est la seule que le miroir JS sait fabriquer."""
    scv = _scv(cfg={"relevance_key": "general", "min_relevance": 0.0})
    doc = cv_docx.build_docx(scv)
    txt = _text(doc)
    assert "ALTEN" in txt and "Projets académiques" not in txt
    assert {p.style.name for p in doc.paragraphs} == {"Normal"}
    assert [p.text for p in doc.paragraphs if p.text.startswith("ALTEN")] == ["ALTEN"]


def test_no_dict_repr_leaks_into_the_document():
    txt = _text(cv_docx.build_docx(_scv("en")))
    assert "{'fr'" not in txt and "{\"fr\"" not in txt


def test_render_bytes_roundtrips_through_python_docx():
    blob = cv_docx.render_docx_bytes(_scv("fr"))
    assert blob[:2] == b"PK"  # conteneur OOXML
    doc = cv_docx.Document(cv_docx.io.BytesIO(blob))
    assert "Robin Denis" in _text(doc)


# ══ M7 — overlay privé (téléphone + phrase de disponibilité) ══════════════════

def test_no_phone_in_the_document_without_an_overlay():
    txt = _text(cv_docx.build_docx(_scv("fr")))
    assert "+33" not in txt


def test_overlay_injects_phone_at_render_time_only():
    """La preuve M7 : le téléphone est DANS le .docx et ABSENT du structured_cv."""
    scv = _scv("fr")
    before = json.dumps(scv, ensure_ascii=False)
    private = {"phone": "+33 6 51 74 81 92",
               "availability": {"fr": "Stage 6 mois — disponible début 2026"}}
    txt = _text(cv_docx.build_docx(scv, private=private))
    assert "+33 6 51 74 81 92" in txt
    assert "Stage 6 mois — disponible début 2026" in txt
    # la projection n'a pas bougé d'un octet
    assert json.dumps(scv, ensure_ascii=False) == before
    assert "phone" not in scv["identity"] and "+33" not in before


def test_overlay_availability_accepts_a_plain_string():
    private = {"availability": "Available from early 2026"}
    txt = _text(cv_docx.build_docx(_scv("en"), private=private))
    assert "Available from early 2026" in txt


def test_overlay_default_path_lives_outside_the_public_repo():
    assert _REPO not in cv_docx.DEFAULT_OVERLAY_PATH.parents


def test_overlay_refuses_a_path_inside_the_public_repo():
    """Garantie STRUCTURELLE : `.gitignore` du dépôt SITE ne couvre pas ce fichier
    (mesuré : `git check-ignore tools/cv/cv_private.json` → exit 1) et un
    `git add -A` publierait le téléphone sur robin-denis.com."""
    for inside in (_REPO / "tools" / "cv" / "cv_private.json", _REPO / "cv_private.json"):
        with pytest.raises(ValueError, match="dépôt public"):
            cv_docx.load_private_overlay(path=inside)


@pytest.mark.parametrize("git_marker", ["dir", "file"])
def test_overlay_refuses_any_git_working_tree_not_just_this_checkout(tmp_path, git_marker):
    """Le dépôt SITE existe en plusieurs checkouts (arbre principal + worktrees) :
    ne garder que `REPO_ROOT` laisserait passer un chemin vers un autre clone.
    `.git` est un RÉPERTOIRE dans un clone et un FICHIER dans un worktree — les
    deux doivent être reconnus (celui de ce worktree-ci fait 79 octets, mesuré)."""
    clone = tmp_path / "autre-checkout"
    (clone / "sous" / "dossier").mkdir(parents=True)
    if git_marker == "dir":
        (clone / ".git").mkdir()
    else:
        (clone / ".git").write_text("gitdir: ../ailleurs\n", encoding="utf-8")
    with pytest.raises(ValueError, match="dépôt public"):
        cv_docx.load_private_overlay(path=clone / "sous" / "dossier" / "cv_private.json")


def test_overlay_refuses_a_path_climbing_back_into_the_repo(tmp_path):
    """`..` lexical : un chemin qui ressort puis rentre doit être attrapé aussi."""
    sneaky = _REPO / "tools" / ".." / "cv_private.json"
    with pytest.raises(ValueError, match="dépôt public"):
        cv_docx.load_private_overlay(path=sneaky)


def test_overlay_fails_loud_on_malformed_json(tmp_path):
    bad = tmp_path / "cv_private.json"
    bad.write_text("{ not json", encoding="utf-8")
    with pytest.raises(ValueError, match="illisible"):
        cv_docx.load_private_overlay(path=bad)


def test_overlay_reads_the_file_and_env_overrides_it(tmp_path):
    src = tmp_path / "cv_private.json"
    src.write_text(json.dumps({"phone": "+33 1", "availability": {"fr": "A", "en": "B"}}),
                   encoding="utf-8")
    got = cv_docx.load_private_overlay(path=src, env={})
    assert got == {"phone": "+33 1", "availability": {"fr": "A", "en": "B"}}
    over = cv_docx.load_private_overlay(path=src, env={"CV_PRIVATE_PHONE": "+33 2"})
    assert over["phone"] == "+33 2"


def test_overlay_is_empty_when_nothing_is_configured(tmp_path):
    got = cv_docx.load_private_overlay(path=tmp_path / "absent.json", env={})
    assert got == {"phone": "", "availability": {"fr": "", "en": ""}}


def test_overlay_env_var_selects_the_file(tmp_path):
    src = tmp_path / "elsewhere.json"
    src.write_text(json.dumps({"phone": "+33 9"}), encoding="utf-8")
    got = cv_docx.load_private_overlay(env={"CV_PRIVATE_OVERLAY": str(src)})
    assert got["phone"] == "+33 9"


def test_overlay_expands_a_tilde_in_the_configured_path(tmp_path, monkeypatch):
    """Le défaut documenté est `~/.elysium/cv_private.json` : le `~` DOIT être résolu.

    Mesuré : retirer `expanduser()` laissait la suite verte, et un `~` dans
    `$CV_PRIVATE_OVERLAY` devenait un répertoire nommé « ~ » sous le cwd — donc un
    fichier absent, donc un overlay silencieusement VIDE (le CV part sans téléphone
    ni disponibilité, sans un mot).
    """
    home = tmp_path / "maison"
    (home / ".elysium").mkdir(parents=True)
    (home / ".elysium" / "cv_private.json").write_text(
        json.dumps({"availability": "Sentinelle-tilde"}), encoding="utf-8")
    for var in ("HOME", "USERPROFILE"):
        monkeypatch.setenv(var, str(home))
    monkeypatch.setattr(pathlib.Path, "home", classmethod(lambda cls: home))
    got = cv_docx.load_private_overlay(
        env={"CV_PRIVATE_OVERLAY": "~/.elysium/cv_private.json"})
    assert got["availability"]["fr"] == "Sentinelle-tilde"


# ══ CLI — `main()` et `_load_cfg()` ═══════════════════════════════════════════
# Mesuré avant ce lot : couverture 83 %, lignes 420-425 (`_load_cfg`) et 431-473
# (`main` ENTIÈRE) jamais exécutées. Deux mutations le prouvaient : `todo = []`
# (le signal « ce qui reste à rédiger » disparaît) et la suppression de
# `stream.reconfigure(encoding="utf-8")` (le correctif cp1252) laissaient toutes
# deux la suite verte. La voie nominale de l'outil n'était pas testée du tout.

_PRIVATE_ENV = ("CV_PRIVATE_OVERLAY", "CV_PRIVATE_PHONE",
                "CV_PRIVATE_AVAILABILITY_FR", "CV_PRIVATE_AVAILABILITY_EN")


@pytest.fixture()
def _cli(tmp_path, monkeypatch):
    """Contexte CLI DÉTERMINISTE : profil contrôlé + aucune variable privée héritée.

    Deux raisons de ne pas laisser `main()` lire le VRAI profile.json ici :
      1. l'overlay privé vit hors dépôt — un poste qui l'a configuré ferait varier
         la sortie, donc le vert dépendrait de la machine ;
      2. surtout, les 10 projets du vrai profil sont AUJOURD'HUI tous sans puce :
         « REDACTION RESTANTE » y est trivialement présent. Le jour où Robin les
         rédige, `todo` devient légitimement vide et le garde deviendrait VACUEUX
         en silence — exactement la panne qu'on répare dans ce lot. Ici un projet
         PORTE des puces et l'autre non : le message doit nommer le second et pas
         le premier, ce qui reste vrai quoi qu'il arrive au vrai profil.
    """
    prof = _profile()
    prof["projects"][0]["bullets"] = {"fr": ["Modélisé le portefeuille de couverture"],
                                      "en": ["Modeled the hedging book"]}
    path = tmp_path / "profile_controle.json"
    path.write_text(json.dumps(prof, ensure_ascii=False), encoding="utf-8")
    monkeypatch.setattr(cv_docx, "_PROFILE_PATH", path)
    for name in _PRIVATE_ENV:
        monkeypatch.delenv(name, raising=False)
    return prof


def test_load_cfg_reads_the_bank_and_fails_loud_on_an_unknown_profile():
    assert cv_docx._load_cfg("risk")["id"] == "risk"        # ni le 1er ni le dernier
    assert cv_docx._load_cfg("full")["relevance_key"] == "general"
    with pytest.raises(SystemExit, match="profil inconnu"):
        cv_docx._load_cfg("nexiste-pas")


def test_main_writes_the_docx_and_names_exactly_what_remains_to_be_written(
        tmp_path, capsys, _cli):
    """La voie nominale de bout en bout — y compris le SIGNAL éditorial sur stderr.

    Muter `todo = cv_select.projects_needing_bullets(...)` en `todo = []` supprime
    la seule chose qui dise à Robin ce qu'il lui reste à rédiger ; rien ne le
    voyait. Le signal est vérifié SÉLECTIF : `pfe-hedging` a des puces et ne doit
    PAS être nommé, `sudoku-cnn` n'en a pas et doit l'être — une chaîne constante
    ne satisferait pas les deux.
    """
    out = tmp_path / "sous" / "dossier" / "cv_full_fr.docx"
    rc = cv_docx.main(["--profile", "full", "--lang", "fr", "--out", str(out),
                       "--overlay", str(tmp_path / "absent.json")])
    assert rc == 0
    assert out.read_bytes()[:2] == b"PK"                    # conteneur OOXML
    doc = cv_docx.Document(str(out))
    assert len(doc.tables) == 0 and {p.style.name for p in doc.paragraphs} == {"Normal"}

    cap = capsys.readouterr()
    assert str(out) in cap.out and "paragraphes" in cap.out
    # Le compte imprimé est celui du FICHIER ÉCRIT — recompté ici depuis le disque.
    # Mesuré : `Document(str(out_path))` → `Document()` faisait imprimer « 0
    # paragraphes » sans que rien ne le voie.
    assert len(doc.paragraphs) > 10, len(doc.paragraphs)
    assert f"{len(doc.paragraphs)} paragraphes" in cap.out, cap.out
    assert f"{len(doc.tables)} table(s)" in cap.out, cap.out
    assert "[cv-docx] overlay prive" in cap.out
    assert "REDACTION RESTANTE" in cap.err, cap.err
    assert "1 projet(s) sans puce : sudoku-cnn" in cap.err, cap.err
    assert "pfe-hedging" not in cap.err, cap.err
    # …et la puce rédigée est bien DANS le document produit
    assert "• Modélisé le portefeuille de couverture" in cv_docx.document_text(doc)


def test_main_survives_a_cp1252_console(tmp_path, monkeypatch, _cli):
    """Leçon M1 (`cv_db_filter.py:319`) : un `print` non ASCII plante en cp1252 et rend
    la voie nominale MORTE — d'où le `stream.reconfigure(encoding="utf-8")` en tête
    de `main`. Supprimer ces lignes laissait la suite verte : rien ne mesurait
    l'encodage de sortie.

    Précision MESURÉE en mutant : le tiret cadratin « — » du message de `main` a bien
    un point de code cp1252 (0x97) — il ne suffit donc PAS à faire tomber la CLI. Ce
    qui la tue, c'est tout caractère absent de cp1252 ; et `main` imprime le chemin
    de sortie tel qu'on le lui donne. On le rend donc explicitement non-cp1252
    (« α », U+03B1). Deux exigences : la CLI ne meurt pas, ET les octets écrits sont
    bien de l'UTF-8 (le tiret cadratin y fait 3 octets, pas 1).
    """
    out = tmp_path / "cv_α.docx"
    raw_out, raw_err = io.BytesIO(), io.BytesIO()
    monkeypatch.setattr(sys, "stdout",
                        io.TextIOWrapper(raw_out, encoding="cp1252", errors="strict"))
    monkeypatch.setattr(sys, "stderr",
                        io.TextIOWrapper(raw_err, encoding="cp1252", errors="strict"))
    rc = cv_docx.main(["--profile", "full", "--lang", "fr", "--out", str(out),
                       "--overlay", str(tmp_path / "absent.json")])
    sys.stdout.flush()
    sys.stderr.flush()
    assert rc == 0
    assert out.exists()
    raw = raw_out.getvalue()
    assert b"\xe2\x80\x94" in raw, "le tiret cadratin n'est pas encodé en UTF-8"
    assert "α" in raw.decode("utf-8"), "le chemin non-cp1252 n'a pas traversé stdout"
    assert "REDACTION RESTANTE" in raw_err.getvalue().decode("utf-8")


def test_main_tolerates_a_stream_without_reconfigure(tmp_path, monkeypatch, _cli):
    """Le correctif cp1252 doit rester OPTIONNEL : un `sys.stdout` sans méthode
    `reconfigure` (StringIO, flux redirigé par un appelant, capture d'un autre
    harnais) ne doit pas tuer la CLI — c'est le rôle du
    `except (AttributeError, ValueError)`, branche jamais exercée jusqu'ici
    (`--cov` : lignes 444-445 manquantes)."""
    monkeypatch.setattr(sys, "stdout", io.StringIO())
    monkeypatch.setattr(sys, "stderr", io.StringIO())
    rc = cv_docx.main(["--profile", "full", "--lang", "fr",
                       "--out", str(tmp_path / "cv.docx"),
                       "--overlay", str(tmp_path / "absent.json")])
    assert rc == 0
    assert "REDACTION RESTANTE" in sys.stderr.getvalue()


def test_main_cli_overrides_the_project_selection(tmp_path, capsys, _cli):
    """`--projects-types` et `--max-projects` sont deux branches de `main` : sans
    test elles pouvaient disparaître sans que rien ne rougisse.

    Le profil contrôlé porte 2 projets non-académiques (`tms-bouygues` 2025,
    `elysium` 2024). Le cap est fixé à 1 — STRICTEMENT sous ce que le filtre laisse
    passer — sinon ignorer `--max-projects` donnerait le même résultat et la
    branche resterait non gardée.
    """
    out = tmp_path / "cv.docx"
    rc = cv_docx.main(["--profile", "full", "--lang", "en", "--out", str(out),
                       "--overlay", str(tmp_path / "absent.json"),
                       "--projects-types", "personal, professional",
                       "--max-projects", "1"])
    assert rc == 0
    cap = capsys.readouterr()
    assert "1 projet(s)" in cap.out, cap.out
    txt = cv_docx.document_text(cv_docx.Document(str(out)))
    assert "TMS" in txt, txt                 # --projects-types a bien été lu
    assert "ELYSIUM" not in txt, txt         # --max-projects a bien coupé à 1
    assert "Sudoku" not in txt, txt          # les académiques sont écartés


def test_main_actually_reads_the_overlay_file_it_is_given(tmp_path, capsys, _cli):
    """Preuve de CÂBLAGE : `--overlay` traverse la CLI jusqu'au document produit.

    Mesuré : remplacer `load_private_overlay(path=args.overlay)` par un dict vide codé
    en dur laissait la suite VERTE — rien ne prouvait que la CLI lise l'overlay privé.
    La sonde est une chaîne sentinelle quelconque (jamais un numéro) : elle ferme le
    trou de câblage sans rien préjuger de l'arbitrage téléphone (ADR-007).
    """
    sentinel = "SENTINELLE-OVERLAY-7F3A"
    overlay = tmp_path / "hors-depot" / "cv_private.json"
    overlay.parent.mkdir(parents=True)
    overlay.write_text(json.dumps({"availability": {"fr": sentinel, "en": sentinel}}),
                       encoding="utf-8")
    out = tmp_path / "cv.docx"
    rc = cv_docx.main(["--profile", "full", "--lang", "fr", "--out", str(out),
                       "--overlay", str(overlay)])
    assert rc == 0
    assert sentinel in cv_docx.document_text(cv_docx.Document(str(out)))
    assert "disponibilite=oui" in capsys.readouterr().out


def test_main_honours_the_requested_language(tmp_path, capsys, _cli):
    """`--lang` doit atteindre la projection.

    Mesuré : câbler `"fr"` en dur à la place de `args.lang` laissait la suite verte —
    `--lang en` produisait un CV français. L'attendu vient de la DONNÉE du profil
    contrôlé (ses champs bilingues), pas d'une chaîne écrite dans le test.
    """
    prof = _cli
    tag, bullet = prof["identity"]["tagline"], prof["experiences"][0]["bullets"]
    got = {}
    for lang in ("fr", "en"):
        out = tmp_path / f"cv_{lang}.docx"
        assert cv_docx.main(["--profile", "full", "--lang", lang, "--out", str(out),
                             "--overlay", str(tmp_path / "absent.json")]) == 0
        got[lang] = cv_docx.document_text(cv_docx.Document(str(out)))
    capsys.readouterr()
    for lang, other in (("fr", "en"), ("en", "fr")):
        assert tag[lang] in got[lang] and tag[other] not in got[lang], lang
        assert bullet[lang][0] in got[lang] and bullet[other][0] not in got[lang], lang


@pytest.mark.skipif(not _REAL_PROFILE.exists(), reason="profile.json réel absent")
def test_main_end_to_end_on_the_real_profile(tmp_path, capsys, monkeypatch):
    """La CLI telle qu'elle est réellement lancée, sur la vraie donnée — l'artefact
    LIVRÉ, pas une reconstitution. Assertions STRUCTURELLES seulement : le contenu
    éditorial du profil bouge, sa forme ATS non."""
    for name in _PRIVATE_ENV:
        monkeypatch.delenv(name, raising=False)
    out = tmp_path / "cv_full_fr.docx"
    rc = cv_docx.main(["--profile", "full", "--lang", "fr", "--out", str(out),
                       "--overlay", str(tmp_path / "absent.json")])
    assert rc == 0
    capsys.readouterr()
    doc = cv_docx.Document(str(out))
    assert len(doc.tables) == 0
    assert {p.style.name for p in doc.paragraphs} == {"Normal"}
    assert 40 <= len(doc.paragraphs) <= 90, len(doc.paragraphs)
    _assert_no_right_column_echo(doc, "CLI/profil réel")


# ══ Smoke sur le VRAI profile.json ════════════════════════════════════════════

@pytest.mark.skipif(not _REAL_PROFILE.exists(), reason="profile.json réel absent")
def test_real_profile_renders_a_structurally_ats_document():
    prof = json.loads(_REAL_PROFILE.read_text(encoding="utf-8"))
    cfg = {"relevance_key": "general", "min_relevance": 0.0, "sort_by": "date",
           "ats_extras": True}
    exps = cv_select.select_experiences(prof, cfg)
    scv = cv_select.build_structured_cv(prof, exps, "fr", cfg)
    doc = cv_docx.build_docx(scv)
    assert len(doc.tables) == 0
    assert {p.style.name for p in doc.paragraphs} == {"Normal"}
    txt = _text(doc)
    assert "Robin Denis" in txt and "ALTEN" in txt and "ECE Paris" in txt
    # les 10 projets académiques du profil réel n'ont AUCUN bullet : la projection
    # doit tous les signaler (travail de rédaction restant, pas de fabrication).
    assert len(scv["projects"]) == 10
    assert cv_select.projects_needing_bullets(scv["projects"]) == [p["id"] for p in scv["projects"]]
    # bande de taille du gabarit mesuré (51-61 paragraphes)
    assert 40 <= len(doc.paragraphs) <= 90, len(doc.paragraphs)
    # …et surtout : le document porte EXACTEMENT ce que la projection lui donne, sur
    # la VRAIE donnée. C'est ici que se voit une section qui disparaît en silence.
    assert [p.text for p in doc.paragraphs] == _expected_outline(scv)
    # 8 des 10 projets portent `date == org == 'ECE Paris'` (quirk de donnée
    # connu et documenté) : le rendu ne doit jamais imprimer deux fois la chaîne.
    doublons = [p for p in scv["projects"]
                if p["org"] and p["org"].strip().casefold() == p["date"].strip().casefold()]
    assert len(doublons) == 8, [p["id"] for p in doublons]   # la mesure, pas un souvenir
    _assert_no_right_column_echo(doc, "profil réel/fr")
    _assert_no_right_column_echo(
        cv_docx.build_docx(cv_select.build_structured_cv(prof, exps, "en", cfg)),
        "profil réel/en")
