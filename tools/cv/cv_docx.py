"""Σ-CV-ATELIER — renderer `.docx` au gabarit ATS personnel (M5/M6/M7, ADR-002).

Entrée : `structured_cv` (la projection déjà produite par `cv_select`), avec le
drapeau `cfg["ats_extras"] = True`. Sortie : un `.docx` structurellement conforme
au gabarit de référence du dépôt.

Ce n'est PAS un export du HTML. Le template `ats` de la banque est CSS-only : il
rend une structure identique à `sobre`. Le gabarit ATS est une STRUCTURE, mesurée
sur les artefacts du dépôt et reproduite ici :

  `cv/template/Template_Resume.docx`      61 paragraphes, 0 table, style `Normal`
                                          seul, titres de section en petites
                                          capitales grasses (10 pt), entrées
                                          alignées par TABULATIONS avec un taquet
                                          DROIT à 6570980 EMU, corps 9 pt,
                                          police Times New Roman, page 8.5x11 in.
  `cv/ATS_FR_RobinDENIS_8.pdf` (+ 2)      mediabox 612x792 pt = US Letter, et la
                                          ligne de contact réellement envoyée :
                                          « Ville • email • téléphone • linkedin ».

Deux écarts DÉLIBÉRÉS par rapport au fichier de référence, mesurés :
  1. Les puces du gabarit sont des listes Word (`numPr`) — invisibles à
     l'extraction de texte (`ref.paragraphs[18].text == 'Executed'`, sans
     marqueur). On écrit une puce LITTÉRALE « • » avec retrait négatif : elle
     survit à tous les extracteurs, ce qui est précisément l'enjeu « ATS-safe ».
     Bonus : `numPr` exigerait une définition de numérotation, donc un style
     autre que `Normal` — le gabarit impose `Normal` seul.
  2. Les dates sont reprises telles que la projection les produit
     (« 2026-02 → présent ») et non reformatées en « Février – Août 2026 » :
     reformater est une décision éditoriale/localisée, pas une projection.

M7 — overlay privé : `identity.phone` est vide dans `profile.json` et
`test_phone_never_projected` interdit au téléphone de traverser
`build_structured_cv` (le dépôt SITE est PUBLIC). Le téléphone et la phrase de
disponibilité s'injectent donc ICI, au seul point de rendu `.docx`, depuis une
source qui vit HORS du dépôt (cf. `load_private_overlay`).

Dépendance : `python-docx` (mesuré : 1.2.0 sur ce poste). Elle n'est pas encore
déclarée dans `requirements.txt` (fichier hors périmètre de ce lot — cf. J2).

Usage :
    python tools/cv/cv_docx.py --profile quant --lang fr --out cv_quant_fr.docx
"""
from __future__ import annotations

import argparse
import io
import json
import os
import pathlib
import sys
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Emu, Pt

import cv_select

__all__ = [
    "Document", "build_docx", "render_docx_bytes", "document_text",
    "load_private_overlay", "DEFAULT_OVERLAY_PATH", "RIGHT_TAB_EMU", "main",
]

_HERE = pathlib.Path(__file__).resolve().parent          # tools/cv
REPO_ROOT = _HERE.parents[1]                             # racine du dépôt SITE
_PROFILE_PATH = REPO_ROOT / "profile.json"
_CFG_PATH = _HERE / "cv_profiles.json"

# ── Constantes MESURÉES sur cv/template/Template_Resume.docx (ne pas « arrondir »)
PAGE_WIDTH_EMU = 7772400        # 8.5 in — US Letter, comme les 3 PDF envoyés
PAGE_HEIGHT_EMU = 10058400      # 11 in
MARGIN_LEFT_EMU = 548640
MARGIN_RIGHT_EMU = 481330
MARGIN_TOP_EMU = 180340
MARGIN_BOTTOM_EMU = 180340
RIGHT_TAB_EMU = 6570980         # taquet DROIT des lignes d'entrée
BASE_FONT = "Times New Roman"

SZ_NAME = Pt(20)                # 254000 EMU dans la référence
SZ_TAGLINE = Pt(11)
SZ_SECTION = Pt(10)             # 127000 EMU — titres en petites capitales
SZ_BODY = Pt(9)                 # 114300 EMU — corps

BULLET = "• "
_BULLET_INDENT = Pt(10)

_LABELS = {
    "fr": {
        "experience": "Expériences professionnelles",
        "projects": "Projets académiques",
        "education": "Formation",
        "skills": "Compétences et langues",
        "interests": "Centres d'intérêt",
        "certifications": "Certifications",
        "languages": "Langues",
    },
    "en": {
        "experience": "Professional Experience",
        "projects": "Academic Projects",
        "education": "Education",
        "skills": "Language & Skills",
        "interests": "Activities & Interest",
        "certifications": "Certifications",
        "languages": "Languages",
    },
}


# ══ M7 — overlay privé ════════════════════════════════════════════════════════

# Défaut HORS du dépôt : le dépôt SITE est public, son `.gitignore` ne compte que
# 4 entrées (mesuré) et n'en couvrirait pas une de plus posée après coup ; un
# `git add -A` publierait le fichier sur robin-denis.com. Un chemin extérieur rend
# la fuite structurellement impossible plutôt que conventionnellement interdite.
DEFAULT_OVERLAY_PATH = pathlib.Path.home() / ".elysium" / "cv_private.json"

ENV_OVERLAY_PATH = "CV_PRIVATE_OVERLAY"
ENV_PHONE = "CV_PRIVATE_PHONE"
ENV_AVAIL_FR = "CV_PRIVATE_AVAILABILITY_FR"
ENV_AVAIL_EN = "CV_PRIVATE_AVAILABILITY_EN"

_EMPTY_OVERLAY = {"phone": "", "availability": {"fr": "", "en": ""}}


def _assert_outside_public_repo(path: pathlib.Path) -> None:
    """Refuse tout chemin situé dans un arbre de travail git.

    Deux gardes, du plus précis au plus général :
      1. ce dépôt-ci (message explicite) ;
      2. **n'importe quel** arbre git en remontant — le dépôt SITE existe en
         plusieurs checkouts (arbre principal + worktrees) et `REPO_ROOT` seul ne
         couvrirait que celui d'où le script tourne. Un `.git` est un répertoire
         (clone) ou un fichier (worktree) : `.exists()` couvre les deux.
    Mesuré : `git check-ignore tools/cv/cv_private.json` sort 1 — rien ne
    l'exclurait d'un `git add -A`, et le dépôt SITE est servi sur robin-denis.com.
    """
    where = None
    if path == REPO_ROOT or REPO_ROOT in path.parents:
        where = REPO_ROOT
    else:
        for anc in path.parents:
            if (anc / ".git").exists():
                where = anc
                break
    if where is not None:
        raise ValueError(
            f"overlay privé refusé : {path} est DANS un dépôt public ({where}). "
            f"Le téléphone n'a pas sa place dans un dépôt servi sur GitHub Pages — "
            f"utilisez {DEFAULT_OVERLAY_PATH} ou la variable {ENV_OVERLAY_PATH}."
        )


def load_private_overlay(path: Any = None, env: dict[str, str] | None = None) -> dict[str, Any]:
    """Charge le complément privé du header ATS (téléphone, disponibilité).

    Résolution : `path` explicite > `$CV_PRIVATE_OVERLAY` > `~/.elysium/cv_private.json`.
    Les variables `CV_PRIVATE_PHONE` / `CV_PRIVATE_AVAILABILITY_{FR,EN}` écrasent le
    fichier (pratique en CI/one-shot, jamais persistées sur disque).

    Fichier absent → overlay vide (le `.docx` se rend sans téléphone). Fichier
    présent mais illisible → ValueError : un overlay silencieusement ignoré
    produirait un CV amputé sans que personne ne le voie (Zero Masking).
    Format attendu : {"phone": "...", "availability": "..." | {"fr": "...", "en": "..."}}
    """
    env = os.environ if env is None else env
    raw = path or env.get(ENV_OVERLAY_PATH) or DEFAULT_OVERLAY_PATH
    p = pathlib.Path(raw).expanduser()
    p = p if p.is_absolute() else (pathlib.Path.cwd() / p)
    _assert_outside_public_repo(pathlib.Path(os.path.normpath(str(p))))

    data: dict[str, Any] = {}
    if p.exists():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, UnicodeDecodeError) as exc:
            raise ValueError(f"overlay privé illisible ({p}) : {exc}") from exc
        if not isinstance(data, dict):
            raise ValueError(f"overlay privé illisible ({p}) : objet JSON attendu, "
                             f"{type(data).__name__} lu")

    avail = data.get("availability")
    if isinstance(avail, str):
        avail = {"fr": avail, "en": avail}
    elif not isinstance(avail, dict):
        avail = {}

    phone = data.get("phone")
    out = {
        "phone": str(env.get(ENV_PHONE) or (phone if isinstance(phone, str) else "") or ""),
        "availability": {
            "fr": str(env.get(ENV_AVAIL_FR) or avail.get("fr") or ""),
            "en": str(env.get(ENV_AVAIL_EN) or avail.get("en") or ""),
        },
    }
    return out


# ══ Primitives de mise en page ════════════════════════════════════════════════

def _para(doc, *, align=None, right_tab: bool = False, bullet: bool = False,
          space_before: Any = None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = space_before
    pf.line_spacing = 1.0
    if align is not None:
        pf.alignment = align
    if right_tab:
        pf.tab_stops.add_tab_stop(Emu(RIGHT_TAB_EMU), WD_TAB_ALIGNMENT.RIGHT)
    if bullet:
        pf.left_indent = _BULLET_INDENT
        pf.first_line_indent = -_BULLET_INDENT
    return p


def _run(p, text: str, *, bold: bool = False, italic: bool = False,
         small_caps: bool = False, size=SZ_BODY):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.small_caps = small_caps
    r.font.size = size
    r.font.name = BASE_FONT
    return r


def _section_title(doc, text: str) -> None:
    """Titre de section : gras + PETITES CAPITALES, 10 pt (mesure de la référence)."""
    p = _para(doc, space_before=Pt(6))
    _run(p, text, bold=True, small_caps=True, size=SZ_SECTION)


def _entry_line(doc, left: str, right: str, *, bold_left: bool = True,
                small_caps_left: bool = True, italic_left: bool = False,
                space_before: Any = None) -> None:
    """Ligne d'entrée : `gauche` + TAB + `droite` alignée sur le taquet DROIT."""
    p = _para(doc, right_tab=True, space_before=space_before)
    _run(p, left, bold=bold_left, small_caps=small_caps_left, italic=italic_left)
    if right:
        _run(p, "\t")
        _run(p, right)


def _bullet_line(doc, text: str) -> None:
    p = _para(doc, bullet=True)
    _run(p, BULLET + text)


def _plain_line(doc, text: str, *, italic: bool = False) -> None:
    p = _para(doc)
    _run(p, text, italic=italic)


def _join(parts, sep: str = " • ") -> str:
    return sep.join(x for x in parts if x)


# ══ M5 — construction du document ═════════════════════════════════════════════

def build_docx(structured_cv: dict[str, Any], private: dict[str, Any] | None = None):
    """Construit le `Document` ATS. `private` = overlay M7 (jamais dans le scv).

    `structured_cv` n'est jamais muté : l'overlay est lu, pas fusionné.
    """
    lang = structured_cv.get("lang") or "fr"
    lab = _LABELS.get(lang, _LABELS["fr"])
    private = private or {}
    phone = str(private.get("phone") or "")
    avail_raw = private.get("availability")
    if isinstance(avail_raw, str):
        availability = avail_raw
    elif isinstance(avail_raw, dict):
        availability = str(avail_raw.get(lang) or "")
    else:
        availability = ""

    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Emu(PAGE_WIDTH_EMU), Emu(PAGE_HEIGHT_EMU)
    s.left_margin, s.right_margin = Emu(MARGIN_LEFT_EMU), Emu(MARGIN_RIGHT_EMU)
    s.top_margin, s.bottom_margin = Emu(MARGIN_TOP_EMU), Emu(MARGIN_BOTTOM_EMU)
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = SZ_BODY
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    # Mesuré : `Document()` de python-docx 1.2.0 naît avec 0 paragraphe — rien à
    # retirer. On le VÉRIFIE plutôt que de le supposer (un gabarit par défaut qui
    # changerait insérerait un paragraphe vide en tête du CV) — et par un `raise`,
    # pas un `assert` : `python -O` supprimerait l'assertion, donc le garde.
    if doc.paragraphs:
        raise RuntimeError(
            f"gabarit python-docx par défaut non vide ({len(doc.paragraphs)} §) : "
            f"le CV commencerait par un paragraphe parasite")

    idy = structured_cv.get("identity") or {}

    # ── Header : nom / sous-titre / disponibilité (privé) / contact ───────────
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, str(idy.get("name") or ""), bold=True, size=SZ_NAME)
    if idy.get("title"):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, str(idy["title"]), size=SZ_TAGLINE)
    if availability:
        # M7 : la phrase de disponibilité vient de l'overlay privé, comme sur le
        # CV EN réellement envoyé (« 6-month internship • Available from early 2026 »).
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, availability, italic=True)
    contact = _join([idy.get("location"), idy.get("email"), phone,
                     idy.get("linkedin"), idy.get("github")])
    if contact:
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, contact)

    # ── Expériences ──────────────────────────────────────────────────────────
    sections = [s_ for s_ in structured_cv.get("sections") or []
                if (s_ or {}).get("kind") == "experience"]
    if sections:
        _section_title(doc, lab["experience"])
        for sec in sections:
            head = _join([sec.get("company"), sec.get("division")], " — ")
            _entry_line(doc, head, str(sec.get("location") or ""), space_before=Pt(3))
            _entry_line(doc, str(sec.get("title") or ""), str(sec.get("dates") or ""),
                        bold_left=True, small_caps_left=False, italic_left=True)
            for b in sec.get("bullets") or []:
                _bullet_line(doc, str(b))

    # ── Projets (M6) ─────────────────────────────────────────────────────────
    projects = structured_cv.get("projects") or []
    if projects:
        _section_title(doc, lab["projects"])
        for pr in projects:
            date_txt = str(pr.get("date") or "")
            org_txt = str(pr.get("org") or "")
            _entry_line(doc, str(pr.get("name") or ""), date_txt, space_before=Pt(3))
            # `project.date` est une chaîne LIBRE dans profile.json, et 8 des 10
            # projets académiques y portent « ECE Paris » — soit exactement
            # l'organisation résolue depuis `context` (mesuré 2026-07-28 sur le vrai
            # profil : atc-simulation, derivatives-pricer, gripper-robot,
            # monte-carlo-gbm, sudoku-cnn, trading-algo-csharp, vba-index-tool,
            # vhdl-calculator). Imprimer les deux met la MÊME chaîne sur deux lignes
            # consécutives du CV envoyé. La donnée n'est pas réécrite ici
            # (profile.json est hors périmètre — le quirk est signalé au rapport) :
            # c'est le RENDU qui refuse de se répéter.
            if org_txt and org_txt.strip().casefold() != date_txt.strip().casefold():
                _entry_line(doc, org_txt, "", bold_left=False,
                            small_caps_left=False, italic_left=True)
            if pr.get("bullets"):
                for b in pr["bullets"]:
                    _bullet_line(doc, str(b))
            elif pr.get("summary"):
                # Repli assumé : la prose existante, jamais une puce fabriquée.
                _plain_line(doc, str(pr["summary"]))
            if pr.get("impact"):
                _plain_line(doc, str(pr["impact"]), italic=True)

    # ── Formation ────────────────────────────────────────────────────────────
    education = structured_cv.get("education") or []
    if education:
        _section_title(doc, lab["education"])
        for e in education:
            _entry_line(doc, str(e.get("school") or ""), str(e.get("period") or ""),
                        space_before=Pt(3))
            # Même règle que le rendu HTML : `title` recouvre souvent `school`
            # (« ECE Paris » vs « ECE Paris, Cycle Ingénieur ») → on l'omet alors.
            title_txt, school_txt = str(e.get("title") or ""), str(e.get("school") or "")
            if title_txt and school_txt and title_txt.startswith(school_txt):
                title_txt = ""
            sub = _join([title_txt, e.get("org")], " — ")
            if sub:
                _plain_line(doc, sub, italic=True)
            if e.get("degree"):
                _plain_line(doc, str(e["degree"]))
            courses = e.get("courses") or []
            if courses:
                _plain_line(doc, f"{e.get('courses_label') or ''} : "
                                 f"{' · '.join(str(c) for c in courses)}".lstrip(" :"))
            cap = e.get("capstone")
            if isinstance(cap, dict):
                cap_txt = _join([cap.get("label"), cap.get("summary")], " — ")
                if cap_txt:
                    _plain_line(doc, cap_txt)

    # ── Compétences / certifications / langues ───────────────────────────────
    groups = structured_cv.get("skills_groups") or []
    certs = structured_cv.get("certifications") or []
    languages = structured_cv.get("languages") or []
    if groups or certs or languages:
        _section_title(doc, lab["skills"])
        for g in groups:
            items = g.get("items") or []
            if items:
                _bullet_line(doc, f"{g.get('label') or ''} : "
                                  f"{' · '.join(str(i) for i in items)}")
        if certs:
            _bullet_line(doc, f"{lab['certifications']} : "
                              f"{' · '.join(str(c) for c in certs)}")
        if languages:
            lang_items = [_join([lg.get("name"), lg.get("level")], " — ") for lg in languages]
            lang_items = [x for x in lang_items if x]
            if lang_items:
                _bullet_line(doc, f"{lab['languages']} : {' · '.join(lang_items)}")

    # ── Centres d'intérêt ────────────────────────────────────────────────────
    interests = structured_cv.get("interests") or []
    if interests:
        _section_title(doc, lab["interests"])
        _bullet_line(doc, " · ".join(str(i) for i in interests))

    return doc


def render_docx_bytes(structured_cv: dict[str, Any],
                      private: dict[str, Any] | None = None) -> bytes:
    buf = io.BytesIO()
    build_docx(structured_cv, private=private).save(buf)
    return buf.getvalue()


def document_text(doc) -> str:
    """Texte extractible du document (ce que voit un ATS)."""
    return "\n".join(p.text for p in doc.paragraphs)


# ══ CLI ═══════════════════════════════════════════════════════════════════════

def _load_cfg(profile_id: str) -> dict[str, Any]:
    cfgs = json.loads(_CFG_PATH.read_text(encoding="utf-8"))["profiles"]
    for c in cfgs:
        if c.get("id") == profile_id:
            return c
    ids = ", ".join(str(c.get("id")) for c in cfgs)
    raise SystemExit(f"[cv-docx] profil inconnu : {profile_id!r} (connus : {ids})")


def main(argv: list[str] | None = None) -> int:
    # Leçon M1 (`cv_db_filter.py:319`) : un `print` non ASCII plante en cp1252 et
    # rend la voie nominale morte. On force l'UTF-8 en sortie avant toute écriture.
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass

    ap = argparse.ArgumentParser(description="Renderer .docx au gabarit ATS personnel")
    ap.add_argument("--profile", default="full", help="id de cv_profiles.json (défaut: full)")
    ap.add_argument("--lang", default="fr", choices=("fr", "en"))
    ap.add_argument("--out", required=True, help="chemin du .docx à écrire")
    ap.add_argument("--overlay", default=None,
                    help=f"overlay privé (défaut: ${ENV_OVERLAY_PATH} ou {DEFAULT_OVERLAY_PATH})")
    ap.add_argument("--projects-types", default=None,
                    help="types de projets, séparés par des virgules (défaut: academic)")
    ap.add_argument("--max-projects", type=int, default=None)
    args = ap.parse_args(argv)

    profile = json.loads(_PROFILE_PATH.read_text(encoding="utf-8"))
    cfg = {**_load_cfg(args.profile), "ats_extras": True}
    if args.projects_types:
        cfg["projects_types"] = [t.strip() for t in args.projects_types.split(",") if t.strip()]
    if args.max_projects is not None:
        cfg["max_projects"] = args.max_projects

    exps = cv_select.select_experiences(profile, cfg)
    scv = cv_select.build_structured_cv(profile, exps, args.lang, cfg)
    private = load_private_overlay(path=args.overlay)

    out_path = pathlib.Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(render_docx_bytes(scv, private=private))

    doc = Document(str(out_path))
    print(f"[cv-docx] {out_path} — {len(doc.paragraphs)} paragraphes, "
          f"{len(doc.tables)} table(s), {len(exps)} experience(s), "
          f"{len(scv.get('projects') or [])} projet(s)")
    print(f"[cv-docx] overlay prive : telephone={'oui' if private['phone'] else 'non'}, "
          f"disponibilite={'oui' if private['availability'].get(args.lang) else 'non'}")
    todo = cv_select.projects_needing_bullets(scv.get("projects") or [])
    if todo:
        print(f"[cv-docx] REDACTION RESTANTE — {len(todo)} projet(s) sans puce : "
              f"{', '.join(todo)}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
